import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import random
import uuid
import os
from docx import Document
from docx.shared import Inches
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.neighbors import KNeighborsRegressor
from sklearn import metrics


async def get_filepath_project() -> str:
    """
    Основная функция для обработки данных, генерации визуализаций 
    и создания итогового отчета в формате Word.
    
    Возвращает:
        str: Путь к сохраненному файлу отчета
    """
    # Инициализация документа и случайных параметров
    doc = Document('data/project.docx')
    random_params = initialize_random_parameters()
    
    # Обновление шаблона документа
    update_document_template(doc, random_params)
    
    # Загрузка и предварительная обработка данных
    df = load_and_preprocess_data()
    
    # Генерация и вставка тепловой карты корреляций
    insert_correlation_heatmap(doc, df, random_params['colour_map'])
    
    # Подготовка данных для моделирования
    X_train, X_test, y_train, y_test = prepare_modeling_data(
        df, 
        random_params['test_size'], 
        random_params['random_state']
    )
    
    # Обновление информации о размерах выборок
    update_dataset_sizes(doc, X_train, X_test)
    
    # Линейная регрессия
    perform_linear_regression(doc, X_train, X_test, y_train, y_test)
    
    # Метод k-ближайших соседей
    perform_knn_regression(doc, X_train, X_test, y_train, y_test)
    
    # Сохранение и возврат итогового документа
    return save_final_document(doc)


def initialize_random_parameters() -> dict:
    """
    Инициализация случайных параметров для анализа.
    
    Возвращает:
        dict: Словарь с параметрами:
            - test_size: размер тестовой выборки (10-35%)
            - random_state: случайное seed-значение
            - colour_map: случайная цветовая схема для визуализаций
    """
    return {
        'test_size': random.randint(10, 35) / 100,
        'random_state': random.randint(0, 150),
        'colour_map': random.choice([
            "viridis", "plasma", "inferno", "magma", "cividis", "spring",
            "summer", "autumn", "winter", "cool", "Wistia", "hot",
            "afmhot", "gist_heat", "copper"
        ])
    }


def update_document_template(doc: Document, params: dict) -> None:
    """
    Замена плейсхолдеров в шаблоне документа на реальные значения.
    
    Аргументы:
        doc: Объект документа Word
        params: Словарь с параметрами для замены
    """
    for paragraph in doc.paragraphs:
        if '{{PROCENT}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                '{{PROCENT}}', 
                str(int(params['test_size'] * 100)))
            
        if '{{RANDOM_STATE}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                '{{RANDOM_STATE}}', 
                str(params['random_state']))
            
        if '{{COLOR}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                '{{COLOR}}', 
                params['colour_map'])


def load_and_preprocess_data() -> pd.DataFrame:
    """
    Загрузка и предварительная обработка данных.
    
    Возвращает:
        pd.DataFrame: Обработанный DataFrame
    """
    df = pd.read_csv('data/ds_salaries.csv', sep=',')
    
    # Удаление ненужных столбцов
    df = df.drop(["salary", "salary_currency"], axis=1)
    
    # Преобразование категориальных переменных в числовые
    mapping_dicts = {
        'experience_level': {'SE': 1, 'MI': 2, 'EN': 3, 'EX': 4},
        'employment_type': {'FT': 1, 'CT': 2, 'FL': 3, 'PT': 4},
        'company_size': {'S': 1, 'M': 2, 'L': 3}
    }
    
    for col, mapping in mapping_dicts.items():
        df[col] = df[col].map(mapping)
    
    return df


def insert_correlation_heatmap(
    doc: Document, 
    df: pd.DataFrame, 
    colour_map: str
) -> None:
    """
    Генерация и вставка тепловой карты корреляций в документ.
    
    Аргументы:
        doc: Объект документа Word
        df: DataFrame с данными
        colour_map: Цветовая схема для визуализации
    """
    corr_matrix = df.corr(numeric_only=True).round(2)
    
    plt.figure(figsize=(9, 6))
    sns.heatmap(corr_matrix, cmap=colour_map, annot=True)
    
    image_path = save_temp_image()
    insert_image_to_doc(doc, '{{IMAGE1}}', image_path)


def prepare_modeling_data(
    df: pd.DataFrame, 
    test_size: float, 
    random_state: int
) -> tuple:
    """
    Подготовка данных для моделирования с разделением на train/test.
    
    Аргументы:
        df: Исходный DataFrame
        test_size: Доля тестовой выборки
        random_state: Seed для воспроизводимости
    
    Возвращает:
        tuple: (X_train, X_test, y_train, y_test)
    """
    feature_columns = df[['work_year', 'experience_level', 'employment_type']]
    target_column = df['salary_in_usd']
    
    return train_test_split(
        feature_columns,
        target_column,
        test_size=test_size,
        random_state=random_state,
    )


def update_dataset_sizes(
    doc: Document, 
    X_train: pd.DataFrame, 
    X_test: pd.DataFrame
) -> None:
    """
    Обновление информации о размерах выборок в документе.
    
    Аргументы:
        doc: Объект документа Word
        X_train: Обучающая выборка
        X_test: Тестовая выборка
    """
    for paragraph in doc.paragraphs:
        if '{{LEANING}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                '{{LEANING}}', 
                str(X_train.shape[0]))
            
        if '{{TEST}}' in paragraph.text:
            paragraph.text = paragraph.text.replace(
                '{{TEST}}', 
                str(X_test.shape[0]))


def perform_linear_regression(
    doc: Document, 
    X_train: pd.DataFrame, 
    X_test: pd.DataFrame, 
    y_train: pd.Series, 
    y_test: pd.Series
) -> None:
    """
    Выполнение линейной регрессии и обновление документа с результатами.
    
    Аргументы:
        doc: Объект документа Word
        X_train: Обучающие признаки
        X_test: Тестовые признаки
        y_train: Обучающая целевая переменная
        y_test: Тестовая целевая переменная
    """
    model = LinearRegression()
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    
    # Расчет метрик
    metrics_data = {
        'ROOT_MEAN1': f'Root Mean Squared Error (RMSE): {np.sqrt(metrics.mean_squared_error(y_test, y_pred))}',
        'R1': f'R2: {np.round(metrics.r2_score(y_test, y_pred), 2)}'
    }
    
    # Обновление документа с метриками
    for key, value in metrics_data.items():
        for paragraph in doc.paragraphs:
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', value)
    
    # Генерация и вставка графика предсказаний
    generate_prediction_plot(y_test, y_pred, "Linear Regression")
    insert_image_to_doc(doc, '{{IMAGE2}}', save_temp_image())


def perform_knn_regression(
    doc: Document, 
    X_train: pd.DataFrame, 
    X_test: pd.DataFrame, 
    y_train: pd.Series, 
    y_test: pd.Series
) -> None:
    """
    Выполнение регрессии k-ближайших соседей и обновление документа.
    
    Аргументы:
        doc: Объект документа Word
        X_train: Обучающие признаки
        X_test: Тестовые признаки
        y_train: Обучающая целевая переменная
        y_test: Тестовая целевая переменная
    """
    model = KNeighborsRegressor()
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    
    # Расчет метрик
    metrics_data = {
        'ROOT_MEAN2': f'Root Mean Squared Error (RMSE): {np.round(np.sqrt(metrics.mean_squared_error(y_test, y_pred)), 2)}',
        'R2': f'R2: {np.round(metrics.r2_score(y_test, y_pred), 2)}'
    }
    
    # Обновление документа с метриками
    for key, value in metrics_data.items():
        for paragraph in doc.paragraphs:
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', value)
    
    # Генерация и вставка графика предсказаний
    generate_prediction_plot(y_test, y_pred, "kNN")
    insert_image_to_doc(doc, '{{IMAGE3}}', save_temp_image())


def generate_prediction_plot(
    y_test: pd.Series, 
    y_pred: np.array, 
    model_name: str
) -> None:
    """
    Генерация графика сравнения предсказанных и реальных значений.
    
    Аргументы:
        y_test: Реальные значения
        y_pred: Предсказанные значения
        model_name: Название модели для легенды
    """
    order = np.argsort(y_test.values)
    y_test_ordered = y_test.values[order]
    y_pred_ordered = y_pred[order]
    
    plt.figure(figsize=(10, 8))
    plt.scatter(y_test_ordered, y_pred_ordered, label=model_name)
    plt.plot(y_test_ordered, y_test_ordered, label="True values", color="red")
    plt.legend()
    plt.xlabel("Истинные значения")
    plt.ylabel("Предсказанные значения")


def save_temp_image() -> str:
    """
    Сохранение текущего графика matplotlib во временный файл.
    
    Возвращает:
        str: Путь к сохраненному изображению
    """
    image_path = f"{uuid.uuid4()}.png"
    plt.savefig(image_path)
    plt.close()
    return image_path


def insert_image_to_doc(
    doc: Document, 
    placeholder: str, 
    image_path: str
) -> None:
    """
    Вставка изображения в документ на место плейсхолдера.
    
    Аргументы:
        doc: Объект документа Word
        placeholder: Плейсхолдер для замены
        image_path: Путь к файлу изображения
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, '')
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(7))
    os.remove(image_path)


def save_final_document(doc: Document) -> str:
    """
    Сохранение итогового документа.
    
    Аргументы:
        doc: Объект документа Word
    
    Возвращает:
        str: Путь к сохраненному файлу
    """
    file_path = f"{uuid.uuid4()}.docx"
    doc.save(file_path)
    return file_path